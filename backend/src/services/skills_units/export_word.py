from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from src.infrastructure.llm.llm_client import LlmClient


class MinutesExportWordSkill:
    def __init__(self, llm_client: LlmClient | None = None) -> None:
        self.llm_client = llm_client

    def run(self, payload: dict[str, Any]) -> dict[str, Any]:
        # 最終確定した議事録データを Word 形式で保存する。
        output_dir = Path(payload["output_dir"])
        output_dir.mkdir(parents=True, exist_ok=True)
        summary_title = self._build_summary_title(payload["final_minutes"], str(payload.get("transcript") or ""))
        today_ymd = datetime.now().strftime("%Y%m%d")
        out_path = output_dir / f"{today_ymd}_{summary_title}.docx"

        doc = Document()
        doc.add_heading("議事録", level=0)
        for section, value in payload["final_minutes"].items():
            doc.add_heading(str(section), level=1)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                doc.add_paragraph(str(value))
        doc.save(out_path)
        return {"artifact_path": str(out_path)}

    def _build_summary_title(self, final_minutes: dict[str, Any], transcript: str) -> str:
        llm_title = self._generate_summary_title_with_llm(final_minutes, transcript)
        if llm_title:
            return llm_title

        # LLM が利用できない場合は、内容から簡易的に要約名を生成する。
        fallback_source = self._extract_fallback_source(final_minutes, transcript)
        return self._normalize_for_filename(fallback_source) or "会議要約"

    def _generate_summary_title_with_llm(
        self, final_minutes: dict[str, Any], transcript: str
    ) -> str | None:
        if not self.llm_client:
            return None
        prompt = "\n".join(
            [
                "あなたは会議資料のファイル名を作るアシスタントです。",
                "議事録内容を要約した、15文字以内の日本語タイトルを1つ返してください。",
                "返答は必ずJSONのみで、次の形式を厳守してください。",
                '{"title":"新製品企画定例"}',
                "記号は使わず、名詞中心で簡潔にしてください。",
                f"議事録: {final_minutes}",
                f"文字起こし: {transcript}",
            ]
        )
        generated = self.llm_client.generate_json(prompt)
        if not isinstance(generated, dict):
            return None
        raw_title = str(generated.get("title") or "").strip()
        normalized = self._normalize_for_filename(raw_title)
        return normalized or None

    def _extract_fallback_source(self, final_minutes: dict[str, Any], transcript: str) -> str:
        if transcript.strip():
            return transcript.strip().splitlines()[0][:30]
        for value in final_minutes.values():
            if isinstance(value, list) and value:
                return str(value[0])[:30]
            normalized = str(value).strip()
            if normalized:
                return normalized[:30]
        return ""

    def _normalize_for_filename(self, raw_text: str) -> str:
        # Windows でも扱えるよう、禁止文字を除去して空白を単一アンダースコアへ寄せる。
        cleaned = re.sub(r"[\\/:*?\"<>|]", "", raw_text).strip()
        cleaned = re.sub(r"\s+", "_", cleaned)
        cleaned = cleaned.strip("._")
        return cleaned[:40]
