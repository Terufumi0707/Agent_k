from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from docx import Document


class Exporter(ABC):
    @abstractmethod
    def export(self, job_id: str, structured_minutes: dict[str, Any], output_dir: Path) -> Path:
        raise NotImplementedError


class DocxExporter(Exporter):
    def export(self, job_id: str, structured_minutes: dict[str, Any], output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / f"{job_id}.docx"
        document = Document()
        document.add_heading("議事録", level=0)

        for section, values in structured_minutes.items():
            document.add_heading(section, level=1)
            if isinstance(values, list):
                if values and isinstance(values[0], dict):
                    table = document.add_table(rows=1, cols=len(values[0]))
                    for i, key in enumerate(values[0].keys()):
                        table.rows[0].cells[i].text = str(key)
                    for row in values:
                        cells = table.add_row().cells
                        for i, value in enumerate(row.values()):
                            cells[i].text = str(value)
                else:
                    for item in values:
                        document.add_paragraph(str(item), style="List Bullet")
            else:
                document.add_paragraph(str(values))

        document.save(file_path)
        return file_path
