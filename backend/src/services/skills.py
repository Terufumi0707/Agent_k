from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document

from src.infrastructure.llm.llm_client import LlmClient


class MinutesTranscribeSkill:
    def run(self, payload: dict[str, Any]) -> dict[str, Any]:
        # 実運用では音声認識結果を返す想定。現在はファイル名を含むダミー文字列を返す。
        audio_path = payload["audio_path"]
        filename = Path(audio_path).name
        return {"transcript": f"[whisper transcript:{filename}] 会議の要点、決定事項、担当タスク。"}


class MinutesDraftSkill:
    def __init__(self, llm_client: LlmClient) -> None:
        self.llm_client = llm_client

    def run(self, payload: dict[str, Any]) -> dict[str, Any]:
        # 文字起こし結果と社内フォーマットを使って、複数の候補案を生成する。
        transcript = payload["transcript"]
        if not transcript.strip():
            raise ValueError("transcript is empty")

        format_definition = payload["company_format"]
        num_candidates = min(max(payload.get("candidate_count", 3), 2), 3)
        feedback = payload.get("feedback")
        base_draft = payload.get("base_draft")
        required_sections = [
            s["name"] for s in format_definition.get("sections", []) if s.get("required", False)
        ]

        prompt = "\n".join(
            [
                "あなたは会議の議事録作成アシスタントです。",
                "入力された文字起こしを基に、必ずJSONのみを返してください。",
                '返却形式は {"candidates": [<議事録オブジェクト>, ...]} です。',
                f"候補数は必ず {num_candidates} 件にしてください。",
                f"必須セクション: {required_sections}",
                f"format_definition={json.dumps(format_definition, ensure_ascii=False)}",
                f"transcript={transcript}",
                f"feedback={feedback or ''}",
                f"base_draft={json.dumps(base_draft, ensure_ascii=False) if base_draft else '{}'}",
            ]
        )

        parsed = self.llm_client.generate_json(prompt)
        candidates = self._normalize_candidates(
            parsed=parsed,
            format_definition=format_definition,
            num_candidates=num_candidates,
            feedback=feedback,
            base_draft=base_draft,
            transcript=transcript,
        )
        return {"candidates": candidates}

    def _normalize_candidates(
        self,
        parsed: dict[str, Any] | None,
        format_definition: dict[str, Any],
        num_candidates: int,
        feedback: str | None,
        base_draft: dict[str, Any] | None,
        transcript: str,
    ) -> list[dict[str, Any]]:
        sections = [s["name"] for s in format_definition.get("sections", [])]
        generated = parsed.get("candidates", []) if parsed else []

        normalized: list[dict[str, Any]] = []
        if isinstance(generated, list):
            for candidate in generated[:num_candidates]:
                if not isinstance(candidate, dict):
                    continue
                for section in sections:
                    candidate.setdefault(section, "")
                normalized.append(candidate)

        if normalized:
            return normalized

        fallback_candidates = []
        for idx in range(1, num_candidates + 1):
            content = {
                "会議概要": f"候補{idx}: {transcript[:60]}",
                "議題": [f"議題{idx}-1", f"議題{idx}-2"],
                "決定事項": [f"決定事項{idx}"],
                "ToDo": [f"担当A: フォローアップ{idx}"],
            }
            for section in sections:
                content.setdefault(section, f"{section}（候補{idx}）")
            if feedback:
                content["レビュー反映"] = feedback
            if base_draft:
                content["元ドラフト参照"] = "あり"
            fallback_candidates.append(content)

        return fallback_candidates


class MinutesReviewSkill:
    def run(self, payload: dict[str, Any]) -> dict[str, Any]:
        # レビュー対象の候補と指示内容を受け取り、承認/差し戻しを判定する。
        candidates = payload["candidates"]
        selected_index = payload["selected_index"]
        instruction = (payload.get("instruction") or "").strip()
        base = candidates[selected_index]
        if instruction and instruction.lower() != "approve":
            # 差し戻し時は元候補をコピーし、レビュー内容を追記した改訂版を返す。
            revised = dict(base)
            revised["レビュー反映"] = instruction
            return {"approved": False, "revised_candidate": revised}
        return {"approved": True, "final_minutes": base}


class MinutesExportWordSkill:
    def run(self, payload: dict[str, Any]) -> dict[str, Any]:
        # 最終確定した議事録データを Word 形式で保存する。
        job_id = payload["job_id"]
        output_dir = Path(payload["output_dir"])
        output_dir.mkdir(parents=True, exist_ok=True)
        out_path = output_dir / f"{job_id}.docx"

        doc = Document()
        doc.add_heading("議事録", level=0)
        for section, value in payload["final_minutes"].items():
            doc.add_heading(str(section), level=1)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                doc.add_paragraph(str(value))
        doc.save(out_path)
        return {"artifact_path": str(out_path)}
